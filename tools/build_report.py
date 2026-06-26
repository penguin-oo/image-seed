from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "doc" / "摄影创作研究助手-项目说明.docx"
DEPLOY_URL = "https://penguin-oo.github.io/image-seed/"
GITHUB_URL = "https://github.com/penguin-oo/image-seed"

SCREENS = [
    ("登录页", ROOT / "output/screenshots/01-login.png"),
    ("首页", ROOT / "output/screenshots/02-home.png"),
    ("AI分析结果页", ROOT / "output/screenshots/03-analysis.png"),
    ("摄影师推荐页", ROOT / "output/screenshots/04-photographers.png"),
    ("摄影书推荐页", ROOT / "output/screenshots/05-books.png"),
    ("收藏页", ROOT / "output/screenshots/06-favorites.png"),
    ("历史记录页", ROOT / "output/screenshots/07-history.png"),
]


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    set_font(run, size=9.5, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def para(doc, text=""):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_font(run)


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, header in enumerate(headers):
        cell_text(tbl.rows[0].cells[i], header, bold=True, color="1F3A5F")
        shade(tbl.rows[0].cells[i], "E8EEF5")
        if widths:
            tbl.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return tbl


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("摄影创作研究助手（Image Seed）")
    set_font(run, size=24, bold=True, color="1F3A5F")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("人工智能模型应用实践课程项目说明文档")
    set_font(run, size=12, color="555555")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("多字段摄影策划器 + MySQL 数据库 + Coze 自带大模型节点工作流")
    set_font(run, size=10.5, color="555555")

    home = ROOT / "output/screenshots/02-home.png"
    if home.exists():
        doc.add_picture(str(home), width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    link = doc.add_paragraph()
    link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = link.add_run("部署链接：" + DEPLOY_URL)
    set_font(run, size=10, color="7A5A00")
    doc.add_page_break()


def build():
    doc = Document()
    setup_styles(doc)
    add_cover(doc)

    doc.add_heading("一、项目简介", level=1)
    para(doc, "Image Seed 是一个面向摄影创作前期研究的 Web 应用。用户填写摄影主题、拍摄地点、拍摄对象、器材、影像风格、周期和交付形式后，系统生成摄影项目策划方案，并从摄影师库、摄影书库中推荐参考对象，帮助学生完成从主题选择到资料研究、拍摄计划和作品呈现的前期准备。")
    para(doc, "本课程作业版不在网页中调用真实外部 API，也不保存 API Key。AI 分析部分使用本地规则模拟大模型输出结构；Coze 部分按“扣子自带大模型节点 + 数据库查询 + 结果整合”的方式提供完整工作流设计和配置。")
    table(doc, ["项目项", "说明"], [
        ["项目名称", "摄影创作研究助手（Image Seed）"],
        ["运行形式", "Web 单页应用，可通过 GitHub Pages 公网访问"],
        ["数据库", "MySQL 8.0 设计文件，含摄影师、摄影书、收藏、历史记录四张表"],
        ["AI/工作流", "本地模拟 AI 输出；Coze 使用扣子自带大模型节点，配置用于课程展示和后续平台搭建"],
        ["部署链接", DEPLOY_URL],
        ["源码仓库", GITHUB_URL],
    ], widths=[1.5, 4.8])

    doc.add_heading("二、设计目的", level=1)
    bullets(doc, [
        "把摄影创作前期的主题分析、参考资料查找、拍摄建议、镜头清单和执行风险整合到一个界面中。",
        "满足课程对 AI 大模型、数据库、工作流和 UI 界面的综合要求。",
        "避免纯后台页面，使用真实摄影素材、明确导航和视觉化工作流，让作品更像可使用的摄影研究工具。",
        "在不调用真实外部 API 的前提下，保留可迁移到 Coze 自带大模型节点的结构化输入、节点和输出格式。",
    ])

    doc.add_heading("三、系统架构", level=1)
    para(doc, "系统采用静态 Web 前端 + 本地数据逻辑 + MySQL 建表文件 + Coze 工作流设计的结构。线上演示由 GitHub Pages 托管，浏览器运行时只加载项目内的 HTML、CSS、JavaScript 和本地图片资源，不依赖本地数据库服务即可演示。")
    table(doc, ["层级", "主要文件", "作用"], [
        ["UI 展示层", "index.html, styles.css, src/app.js", "登录、首页、AI 分析、摄影师推荐、摄影书推荐、收藏、历史记录页面"],
        ["业务逻辑层", "src/core.js", "多字段 brief 解析、关键词匹配、摄影师/摄影书推荐、本地 AI 风格输出"],
        ["数据库设计层", "database/image_seed.sql", "MySQL 表结构和 20 条摄影师、20 条摄影书测试数据"],
        ["工作流设计层", "docs/coze_workflow.md, docs/coze_workflow_config.json", "Coze 自带大模型节点、变量、提示词、输出结构说明"],
        ["素材层", "assets/photos, assets/image_sources.json", "真实摄影相关图片素材及来源记录"],
    ], widths=[1.2, 2.2, 2.9])

    doc.add_heading("四、工作流程说明", level=1)
    numbered(doc, [
        "用户在登录页输入用户名进入系统。",
        "用户在首页填写摄影主题、地点、对象、器材、风格、周期和交付形式，或选择预设主题标签。",
        "系统生成 AI 风格创作分析，包括标题、概念、方向、拍摄建议、视觉元素、呈现方式、日程、成片清单、风险和工具包。",
        "系统根据关键词匹配摄影师数据表，展示相关摄影师。",
        "系统根据关键词匹配摄影书数据表，展示相关摄影书。",
        "用户可收藏生成方案、摄影师或摄影书。",
        "系统把每次生成记录写入浏览器 localStorage，模拟历史记录表。",
    ])
    para(doc, "Coze 工作流对应节点：开始节点 -> 扣子自带大模型节点 - 摄影项目策划 -> 摄影师数据库查询节点 -> 摄影书数据库查询节点 -> 结果整合节点 -> 历史记录节点 -> 输出节点。")
    para(doc, "Coze 使用情况：已使用用户提供的 www.coze.cn cookie 访问 Coze，能够进入开放平台 Playground 并查看工作流、数据库相关能力入口；点击“创建”时跳转火山引擎登录页，需要账号二次登录。因此本次交付提供完整 Coze 工作流配置文件，后续登录后可按配置搭建。")

    doc.add_heading("五、数据库设计", level=1)
    table(doc, ["表名", "字段", "说明"], [
        ["photographers", "id, name, country, style, keywords, intro", "摄影师信息表，已内置至少 20 条测试数据"],
        ["photo_books", "id, title, author, theme, intro", "摄影书信息表，已内置至少 20 条测试数据"],
        ["favorites", "id, username, favorite_type, favorite_content, created_at", "收藏表，用于保存用户收藏内容"],
        ["history_records", "id, username, keyword, project_brief, ai_result, created_at", "历史记录表，用于保存生成过的主题、多字段项目 brief 和结果"],
    ], widths=[1.45, 2.9, 2.1])
    para(doc, "数据库文件位置：database/image_seed.sql。该文件可直接导入 MySQL 8.0，包含建库、建表和测试数据插入语句。")

    doc.add_heading("六、功能介绍", level=1)
    table(doc, ["页面", "功能说明"], [
        ["登录页", "简单用户名登录，进入课程演示系统。"],
        ["首页", "填写多字段摄影项目 brief，展示 Coze/AI/数据库工作流程，并提供真实摄影素材图。"],
        ["AI 分析结果页", "展示项目标题、创作概念、创作方向、拍摄建议、视觉元素、呈现方式、拍摄日程、12 张成片清单、风险提醒和现场工具包。"],
        ["摄影师推荐页", "从摄影师数据中按关键词匹配相关摄影师。"],
        ["摄影书推荐页", "从摄影书数据中按关键词匹配相关摄影书。"],
        ["收藏页", "可收藏方案、摄影师或摄影书，使用 localStorage 模拟收藏表。"],
        ["历史记录页", "显示过去生成的主题和时间，使用 localStorage 模拟历史记录表。"],
    ], widths=[1.5, 4.8])

    doc.add_heading("七、页面截图", level=1)
    for title, image in SCREENS:
        if image.exists():
            doc.add_heading(title, level=2)
            doc.add_picture(str(image), width=Inches(6.25))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("八、部署链接", level=1)
    para(doc, "公网访问链接：" + DEPLOY_URL)
    para(doc, "GitHub 源码仓库：" + GITHUB_URL)
    para(doc, "本地运行方式：node server.cjs 8788 127.0.0.1，然后访问 http://127.0.0.1:8788/index.html。")

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Image Seed 摄影创作研究助手 - 课程项目说明")
        set_font(run, size=9, color="777777")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
